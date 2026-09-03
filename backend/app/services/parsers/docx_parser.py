import io
from typing import Tuple, Dict, Any
from docx import Document
from fastapi import HTTPException, status
from app.services.parsers.base import BaseDocumentParser


class DocxDocumentParser(BaseDocumentParser):
    def parse(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to parse DOCX document (malformed or corrupted archive): {str(e)}"
            )

        elements = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                level = 1
                try:
                    level = int(paragraph.style.name.replace("Heading", "").strip())
                except ValueError:
                    pass
                elements.append(f"{'#' * level} {text}")
            else:
                elements.append(text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                elements.append("\n" + "\n".join(table_rows) + "\n")

        full_text = "\n\n".join(elements).strip()
        if not full_text:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="DOCX document contains no extractable text content."
            )

        return full_text, {
            "format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
